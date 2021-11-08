#!/usr/bin/env python3
#
# abbr_finder.py
#
# Tyler W. Davis
#
# Searches a .docx document for abbreviations and acronyms.
#
##############################################################################
# IMPORT NECESSARY MODULES
##############################################################################
import os

import docx


##############################################################################
# FUNCTIONS
##############################################################################
def read_paragraphs(d):
    """
    Returns a dictionary of paragraph-level text information
    """
    paras = {}
    num_paras = len(d.paragraphs)
    for i in range(num_paras):
        para = d.paragraphs[i]
        paras[i] = {
            "text": para.text
        }
    return paras


##############################################################################
# MAIN
##############################################################################
if __name__ == '__main__':
    docx_file = "example-3.docx"
    docx_path = os.path.join("examples", docx_file)
    my_doc = docx.Document(docx_path)
    my_paras = read_paragraphs(my_doc)
    my_keys = sorted(list(my_paras.keys()))
    for k in my_keys:
        print(k, my_paras[k]['text'])
        # TODO: search for abbreviations
        # exclude parentheticals starting with "e.g.", "i.e.", "see", "refer"
        # find both abbreviation definitions and definition abbreviations
