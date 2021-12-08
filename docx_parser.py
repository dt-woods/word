#!/usr/bin/env python3
#
# docx_parser.py
#
# VERSION: 2.0.0
# UPDATED: 2021-11-24
#
##############################################################################
# PUBLIC DOMAIN NOTICE                                                       #
##############################################################################
# This software is freely available to the public for use.                   #
#                                                                            #
# Although all reasonable efforts have been taken to ensure the accuracy and #
# reliability of the software, the author does not and cannot warrant the    #
# performance or results that may be obtained by using this software.        #
# The author disclaims all warranties, express or implied, including         #
# warranties of performance, merchantability or fitness for any particular   #
# purpose.                                                                   #
#                                                                            #
# Please cite the author in any work or product based on this material.      #
#    Tyler W. Davis, PhD                                                     #
#    https://github.com/dt-woods/                                            #
##############################################################################
#
##############################################################################
# REQUIRED MODULES
##############################################################################
import docx

from docx_utils import delete_paragraph
from docx_utils import find_word_files
from docx_utils import list_paragraph_styles


##############################################################################
# FUNCTIONS
##############################################################################
def parse_file(doc, style, idx):
    """
    Name:     parse_file
    Inputs:   - str, file path to .docx (doc)
              - str, the .docx paragraph style ID to break on (style)
              - int, the index of style ID to parse; zero indexed (idx)
    Features: Finds paragraphs of the given style and breaks it into a
              separate document.
    Depends:  delete_paragraph
    """
    d = docx.Document(my_file)
    para_num = len(d.paragraphs)
    j = 0      # track paragraphs with matching styles
    f = False  # track all paragraphs between matching styles
    for para in d.paragraphs:
        if para.style.style_id == style:
            if j == idx:
                f = True
            else:
                f = False
                delete_paragraph(para)
            j += 1
        elif not f:
            delete_paragraph(para)

    # Save the last chapter
    out_name = "DOCUMENT-%d.docx" % (idx)
    d.save(out_name)


##############################################################################
# MAIN
##############################################################################
if __name__ == '__main__':
    # User inputs:
    my_dir = "examples"     # where to look for the input document
    my_key = "example-1"    # keyword for finding the right input document
    br_style = "Heading1" # the paragraph style used to parse the input document

    # Step 1: find the input word file(s)
    my_files = find_word_files(my_dir, my_key)
    if len(my_files) == 1:
        my_file = my_files[0]
    elif len(my_files) > 1:
        print("Found several word files; "
              "please use keywords to specify the one you want.")
    else:
        print("Failed to find docx. Please check and try again.")
        my_file = None

    if my_file:
        # Step 2 - Find all styles and see if break style is there
        my_doc = docx.Document(my_file)
        my_styles = list_paragraph_styles(my_doc)
        if br_style in my_styles.keys():
            # Step 3 - For each break style, parse:
            for i in range(my_styles[br_style]['count']):
                parse_file(my_file, br_style, i)
