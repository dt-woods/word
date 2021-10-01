#!/usr/bin/env python3
#
# docx_merger.py
#
# VERSION: 0.2.0
# UPDATED: 2021-10-01
#
#
##############################################################################
# PUBLIC DOMAIN NOTICE                                                       #
##############################################################################
# This software is freely available to the public for use.                   #
#                                                                            #
# Although all reasonable efforts have been taken to ensure the accuracy and #
# reliability of the software, the author does not and cannot warrant the    #
# performance or results that may be obtained by using this software.        #
# The author disclaims all warranties, express or implied, including         #
# warranties of performance, merchantability or fitness for any particular   #
# purpose.                                                                   #
#                                                                            #
# Please cite the author in any work or product based on this material.      #
#    Tyler W. Davis, PhD                                                     #
#    https://github.com/dt-woods/                                            #
#                                                                            #
##############################################################################
#
#
##############################################################################
# REQUIRED MODULES
##############################################################################
import os

import docx
from docx.enum.section import WD_SECTION

from docx_utils import find_word_files
from docx_utils import match_char_style
from docx_utils import match_sect_properties


##############################################################################
# FUNCTIONS
##############################################################################
def merge_files(d_list, sbreak):
    """
    Name:     merge_files
    Inputs:   - list, Word documents to merge (d_list)
              - docx.enum.base.EnumValue, section break type btn merged docs
    Outputs:  docx.document.Document, merged document
    Features: Concatenates word .docx files together preserving paragraph
              styles, character formatting (e.g., bold) and section properties
              (e.g., page dimensions and margins). Assumes a new page between
              each merged document.
    Depends:  - match_char_style
              - match_sect_properties
    """
    # Initialize emtpy return document
    out_doc = docx.Document()

    # Iterate over each file
    num_files = len(d_list)
    for i in range(num_files):
        my_file = d_list[i]
        my_doc = docx.Document(my_file)

        # Match section properties (assumes input file has only 1 section)
        out_sect = out_doc.sections[i]
        mat_sect = my_doc.sections[0]
        match_sect_properties(mat_sect, out_sect)

        # Iterate over each paragraph and append to new doc
        for para in my_doc.paragraphs:
            # Create a new empty paragraph, then iterate over paragraph runs
            # NOTE: every paragraph has at least one run
            out_p = out_doc.add_paragraph(text="", style=para.style.name)
            for p_run in para.runs:
                out_r = out_p.add_run(
                    text = p_run.text, style = p_run.style.name)
                match_char_style(p_run, out_r)
        # Create a new section for each new merged file (assumes new page)
        if i < num_files - 1:
            out_doc.add_section(sbreak)
    return out_doc

##############################################################################
# MAIN
##############################################################################
# User inputs:
my_dir = "."            # where to look for the input document
my_key = "DOCUMENT_"    # keyword for finding the right input document
sect_break = WD_SECTION.NEW_PAGE   # section break type between merged files
out_file = "{}_ALL.docx".format(my_key)

# Step 1: find the input word files
my_files = find_word_files(my_dir, my_key)
num_files = len(my_files)
if num_files == 0:
    print("Failed to find any files. "
          "Please update path and keywords and try again.")
else:
    cat_doc = merge_files(my_files, sect_break)
    if os.path.isfile(out_file):
        print("Warning: overwriting existing output file!")
    cat_doc.save(out_file)
