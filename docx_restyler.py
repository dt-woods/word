#!/usr/bin/env python3
#
# docx_restyler.py
#
# VERSION: 2.0.0
# UPDATED: 2021-11-07
#
#
##############################################################################
# PUBLIC DOMAIN NOTICE                                                       #
##############################################################################
# This software is freely available to the public for use.                   #
#                                                                            #
# Although all reasonable efforts have been taken to ensure the accuracy and #
# reliability of the software, the author does not and cannot warrant the    #
# performance or results that may be obtained by using this software.        #
# The author disclaims all warranties, express or implied, including         #
# warranties of performance, merchantability or fitness for any particular   #
# purpose.                                                                   #
#                                                                            #
# Please cite the author in any work or product based on this material.      #
#    Tyler W. Davis, PhD                                                     #
#    https://github.com/dt-woods/                                            #
##############################################################################
#
##############################################################################
# REQUIRED MODULES
##############################################################################
import json
import os
import re
import warnings

import docx
from docx.dml.color import ColorFormat
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from docx_utils import find_files
from docx_utils import list_paragraph_styles
from docx_utils import match_char_style
from docx_utils import match_sect_properties


##############################################################################
# FUNCTIONS
##############################################################################
def add_custom_para_style(d, sname, sloc = 'styles'):
    """
    Name:     add_custom_para_style
    Inputs:   - docx.document.Document, docx document object (d)
              - str, unique custom style name (sname)
              - str, directory for style definitions (sloc)
    Outputs:  None
    Features: Add a custom Paragraph style to a docx Document
    Depends:  - make_style
              - read_style
    """
    doc_styles = d.styles
    sdef = read_style(sname, sloc)

    # Check that the style file exists and isn't already defined:
    if sdef and sname not in doc_styles:
        new_style = doc_styles.add_style(
            name = sname,
            style_type = WD_STYLE_TYPE.PARAGRAPH
        )
        make_style(new_style, sdef, doc_styles)
    elif sdef and sname in doc_styles:
        warnings.warn("{} already exists! Skipping.".format(sname))
        return 1
    else:
        return 1


def apply_style(orig_doc, st_map):
    """
    Name:     apply_style
    Inputs:   - docx.document.Document, the original document (orig_doc)
              - dict, style map (st_map)
    Outputs:  None.
    Features: Applies new styles
    Depends:  - match_char_style
              - match_sect_properties
    TODO:     _ Add a section break paragraph style (if one) so as to
                maintain section styles as the doc is being copied over
    """
    num_sect = len(orig_doc.sections)
    num_para = len(orig_doc.paragraphs)
    print("Copying {} sections and {} paragraphs.".format(num_sect, num_para))

    # Iterate over original document
    for para in orig_doc.paragraphs:
        pstyle_name = para.style.name
        pstyle_id = para.style.style_id
        if pstyle_name in st_map.keys():
            new_style = st_map[pstyle_name]
        elif pstyle_id in st_map.keys():
            new_style = st_map[pstyle_id]
        else:
            # If not mapped, use the original
            new_style = pstyle_name

        if new_style in orig_doc.styles:
            para.style = new_style
        else:
            print("Style {} undefined; using original".format(new_style))


def lookup_style(sname, sloc = 'styles'):
    """
    Name:     lookup_style
    Inputs:   - str, custom style name (sname)
              - str, directory for style definitions (sloc)
    Features: Checks a directory for a given style definition file
    """
    r = re.compile(" ")
    style_name = r.sub("_", sname.lower())
    my_styles = find_files(sloc, "{}*".format(style_name))
    num_styles = len(my_styles)
    if num_styles == 0 or num_styles > 1:
        warnings.warn("Found {} style matches!".format(num_styles))
        return None
    else:
        return my_styles[0]


def make_style(s, d, bs):
    """
    Name:     make_style
    Inputs:   - docx.styles.style._ParagraphStyle, style object (s)
              - dict, paragraph style definitions (d)
              - docx.styles.styles.Styles, output file's style definitions (bs)
    Features: Defines style parameters based on JSON definition
    """
    # Grab the font, color and paragraph style objects
    font_style = s.font
    color_style = font_style.color
    para_style = s.paragraph_format

    if 'definition' in d.keys():
        my_def = d['definition']
        my_keys = list(my_def.keys())
        # FONT STYLES
        if 'font' in my_def.keys():
            my_keys.pop(my_keys.index('font'))
            font_def = my_def['font']
            tf_font_keys = ['all_caps', 'bold', 'double_strike', 'emboss',
                       'hidden', 'imprint', 'italic', 'math',
                       'no_proof', 'outline', 'rtl', 'shadow',
                       'small_caps', 'snap_to_grid', 'spec_vanish', 'strike',
                       'subscript', 'superscript', 'underline', 'web_hidden']
            for k, v in font_def.items():
                if k in tf_font_keys:
                    # Execute True/False assignments
                    exec("font_style.{} = {}".format(k, v))
                else:
                    # The other params of interest are color, name and size
                    if k == 'name':
                        font_style.name = v
                    elif k == 'color':
                        if isinstance(v, dict):
                            if 'hex' in v.keys():
                                color_style.rgb = RGBColor.from_string(v['hex'])
                            else:
                                print("Unknown color keys:", v.keys())
                        else:
                            print("Unknown color value:", v)
                    elif k == 'size':
                        if isinstance(v, dict):
                            if 'point' in v.keys():
                                font_style.size = Pt(v['point'])
                            else:
                                print("Unknown size keys:", v.keys())
                        else:
                            print("Unknown font size:", v)
        # PARAGRAPH FORMAT STYLES
        if 'paragraph_format' in my_def.keys():
            my_keys.pop(my_keys.index('paragraph_format'))
            para_def = my_def['paragraph_format']
            tf_pf_keys = ['keep_together', 'keep_with_next',
                          'page_break_before', 'window_control']
            len_pf_keys = ['first_line_indent', 'left_indent',
                           'line_spacing', 'right_indent',
                           'space_before', 'space_after']
            for k, v in para_def.items():
                if k in tf_pf_keys:
                    # Execute True/False assignment
                    exec("para_style.{} = {}".format(k, v))
                elif k in len_pf_keys:
                    # v should be a dict w/ 'point' 'inch' 'cm' or 'value'
                    if isinstance(v, dict):
                        if 'value' in v.keys():
                            # Use raw number
                            exec("para_style.{} = {}".format(k, v['value']))
                        elif 'point' in v.keys():
                            # Create a point object
                            exec("para_style.{} = Pt({})".format(k, v['point']))
                        elif 'inch' in v.keys():
                            exec("para_style.{} = Inches({})".format(
                                k, v['inch']))
                        else:
                            print("Unknown length key:", v.keys())
                    else:
                        print("Unknown length value:", v)
                elif k == 'alignment':
                    if 'v' == 'left':
                        para_style.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif 'v' == 'right':
                        para_style.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif 'v' == 'center':
                        para_style.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif 'v' == 'justify':
                        para_style.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        print("Unknown alignment:", v)
                else:
                    print("Unknown paragraph format key:", k)
        # The remaining paragraph styles:
        tf_para_keys = ['builtin', 'hidden', 'locked', 'quick_style']
        for my_key in my_keys:
            my_val = my_def[my_key]
            if my_key in tf_para_keys or my_key == 'priority':
                # Execute True/False assignment
                exec("s.{} = {}".format(my_key, my_val))
            elif my_key == 'base_style' and my_val in bs:
                s.base_style = bs[my_val]
            elif my_key == 'next_paragraph_style' and my_val in bs:
                s.next_paragraph_style = bs[my_val]
            else:
                # NOTE: apply new next_paragraph_style later
                print("Unused key:", my_key)


def read_style(sname, sloc = 'styles'):
    """
    Name:     read_style
    Inputs:   - str, custom style name
              - str, directory for style definitions (sloc)
    Features: Reads a JSON style file and returns a dictionary object
    Depends:  lookup_style
    """
    sdef = lookup_style(sname, sloc)
    if sdef:
        with open(sdef, 'r') as f:
            data = f.read()
        style = json.loads(data)
    else:
        style = None
    return style


##############################################################################
# MAIN
##############################################################################
# User inputs:
my_dir = "examples"          # where to look for the input document
my_key = "example-1.docx"    # keyword for finding the right input document

# Define the old-to-new style mapping:
style_map = {
    'Heading1': 'New Head1',
    'Normal': "New Normal"
}

my_files = find_files(my_dir, my_key)

for my_file in my_files:
    # Define the output file name and location
    out_file = "{}_styled.docx".format(
        os.path.basename(my_file).split(".docx")[0]
    )
    # Open existing and new empty docx objects
    my_doc = docx.Document(my_file)

    # Add custom styles to new docx object
    for new_style in style_map.values():
        add_custom_para_style(my_doc, new_style)

    # TODO: add custom next_paragraph_styles after all new styles are defined.

    # Copy old file content to new file, but use new styles.
    apply_style(my_doc, style_map)

    # Save the copy
    my_doc.save(out_file)
